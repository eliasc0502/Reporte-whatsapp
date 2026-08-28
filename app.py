"""
Reporte de avance - Campaña WhatsApp a clientes potenciales
-------------------------------------------------------------
Cada semana se sube un Excel con la lista COMPLETA de clientes
(los antiguos + los nuevos agregados). La app compara esa lista
contra el último estado conocido (guardado en SQLite) para
detectar qué cambió ESTA semana, y así no duplicar conteos.

Columnas esperadas en el Excel:
Nombre | Contacto | Email | Localidad | Región | Estado | Ejecutivo | Venta

- Email, Localidad, Región y Venta son opcionales (pueden venir vacías
  o no existir todavía).
- Estado esperado: "Whatsapp enviado", "Sin Whatsapp", "Derivado", o vacío.
"""

import re
import unicodedata
from datetime import datetime
from io import BytesIO

import gspread
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Inches
from google.oauth2.service_account import Credentials

REQUIRED_COLUMNS = ["Nombre", "Contacto", "Estado", "Ejecutivo"]
OPTIONAL_COLUMNS = ["Email", "Localidad", "Región", "Venta"]

ESTADO_ENVIADO = "Whatsapp enviado"
ESTADO_SIN_WHATSAPP = "Sin Whatsapp"
ESTADO_DERIVADO = "Derivado"

MASTER_COLUMNS = [
    "cliente_key", "nombre", "contacto", "email", "localidad", "region",
    "estado", "ejecutivo", "venta", "semana_ultimo_cambio", "fecha_actualizacion",
]
HISTORIAL_COLUMNS = [
    "semana", "whatsapp_enviados", "sin_whatsapp", "derivados",
    "ventas_total", "fecha_carga",
]
LOG_CARGAS_COLUMNS = [
    "fecha_hora", "semana", "nombre_archivo", "filas_en_archivo",
    "enviados_nuevos", "sin_whatsapp_nuevos", "derivados_nuevos", "ventas_detectadas",
]

REGIONES_CHILE = {
    "1": "Tarapacá", "2": "Antofagasta", "3": "Atacama", "4": "Coquimbo",
    "5": "Valparaíso", "6": "O'Higgins", "7": "Maule", "8": "Biobío",
    "9": "Araucanía", "10": "Los Lagos", "11": "Aysén", "12": "Magallanes",
    "13": "Metropolitana", "14": "Los Ríos", "15": "Arica y Parinacota", "16": "Ñuble",
}

SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]


# ---------------------------------------------------------------------------
# Conexión a Google Sheets
# ---------------------------------------------------------------------------

def get_spreadsheet():
    """Se conecta a la Google Sheet usando la cuenta de servicio configurada
    en los 'Secrets' de Streamlit (ver README para cómo configurarlos)."""
    creds_dict = dict(st.secrets["gcp_service_account"])
    creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
    client = gspread.authorize(creds)
    return client.open_by_key(st.secrets["sheet_id"])


def _obtener_hoja(sh, nombre, columnas):
    """Devuelve la pestaña (worksheet) pedida, creándola con su encabezado
    si todavía no existe."""
    try:
        return sh.worksheet(nombre)
    except gspread.exceptions.WorksheetNotFound:
        ws = sh.add_worksheet(title=nombre, rows=1000, cols=max(len(columnas), 5))
        ws.append_row(columnas)
        return ws


def load_master(sh):
    ws = _obtener_hoja(sh, "clientes_maestro", MASTER_COLUMNS)
    registros = ws.get_all_records()
    if not registros:
        return pd.DataFrame(columns=MASTER_COLUMNS)
    df = pd.DataFrame(registros)
    df = df.replace("", pd.NA)
    if "venta" in df.columns:
        df["venta"] = pd.to_numeric(df["venta"], errors="coerce")
    if "semana_ultimo_cambio" in df.columns:
        df["semana_ultimo_cambio"] = pd.to_numeric(df["semana_ultimo_cambio"], errors="coerce")
    return df


def load_historial(sh):
    ws = _obtener_hoja(sh, "historial_semanal", HISTORIAL_COLUMNS)
    registros = ws.get_all_records()
    if not registros:
        return pd.DataFrame(columns=HISTORIAL_COLUMNS)
    df = pd.DataFrame(registros)
    for col in ["semana", "whatsapp_enviados", "sin_whatsapp", "derivados", "ventas_total"]:
        df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)
    return df.sort_values("semana").reset_index(drop=True)


def guardar_master(sh, df_master):
    ws = _obtener_hoja(sh, "clientes_maestro", MASTER_COLUMNS)
    ws.clear()
    if df_master.empty:
        ws.update([MASTER_COLUMNS])
        return
    cuerpo = df_master[MASTER_COLUMNS].fillna("").astype(str).values.tolist()
    ws.update([MASTER_COLUMNS] + cuerpo)


def guardar_historial(sh, df_hist):
    ws = _obtener_hoja(sh, "historial_semanal", HISTORIAL_COLUMNS)
    ws.clear()
    if df_hist.empty:
        ws.update([HISTORIAL_COLUMNS])
        return
    cuerpo = df_hist[HISTORIAL_COLUMNS].fillna("").astype(str).values.tolist()
    ws.update([HISTORIAL_COLUMNS] + cuerpo)


def registrar_carga(sh, semana, nombre_archivo, filas_en_archivo, resultado):
    """Deja registrada cada carga de Excel (fecha, semana, archivo y qué
    contó), para poder rastrear qué se subió y cuándo."""
    ws = _obtener_hoja(sh, "historial_cargas", LOG_CARGAS_COLUMNS)
    fila = [
        datetime.now().strftime("%Y-%m-%d %H:%M"),
        str(semana),
        nombre_archivo,
        str(filas_en_archivo),
        str(resultado["enviados"]),
        str(resultado["sin_whatsapp"]),
        str(resultado["derivados"]),
        str(resultado["ventas_total"]),
    ]
    ws.append_row(fila)


def load_log_cargas(sh):
    ws = _obtener_hoja(sh, "historial_cargas", LOG_CARGAS_COLUMNS)
    registros = ws.get_all_records()
    if not registros:
        return pd.DataFrame(columns=LOG_CARGAS_COLUMNS)
    df = pd.DataFrame(registros)
    return df.iloc[::-1].reset_index(drop=True)  # más reciente primero (orden de carga real)


# ---------------------------------------------------------------------------
# Normalización / claves de cliente
# ---------------------------------------------------------------------------

def normalize_text(s):
    if pd.isna(s):
        return ""
    s = str(s).strip().lower()
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")
    s = re.sub(r"\s+", " ", s)
    return s


def build_key(nombre, contacto):
    digits = re.sub(r"\D", "", str(contacto)) if pd.notna(contacto) else ""
    if len(digits) >= 8:
        return "tel_" + digits[-9:]
    return "name_" + normalize_text(nombre)


# ---------------------------------------------------------------------------
# Validación de archivo subido
# ---------------------------------------------------------------------------

# Mapa de nombre de columna "sin mayúsculas/tildes" -> nombre canónico que
# usa el resto del código. Así el Excel puede llegar con las columnas en
# MAYÚSCULAS, minúsculas, o con/sin tilde en "Región", y siempre funciona.
COLUMNA_CANONICA = {
    "NOMBRE": "Nombre",
    "CONTACTO": "Contacto",
    "EMAIL": "Email",
    "LOCALIDAD": "Localidad",
    "REGION": "Región",
    "ESTADO": "Estado",
    "EJECUTIVO": "Ejecutivo",
    "VENTA": "Venta",
}


def _clave_columna(nombre_col):
    s = str(nombre_col).strip()
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")
    return s.upper()


def normalizar_columnas(df):
    """Renombra las columnas del Excel subido a los nombres canónicos,
    sin importar mayúsculas/minúsculas o tildes."""
    renombres = {}
    for col in df.columns:
        clave = _clave_columna(col)
        if clave in COLUMNA_CANONICA:
            renombres[col] = COLUMNA_CANONICA[clave]
    return df.rename(columns=renombres)


def validar_columnas(df):
    faltantes = [c for c in REQUIRED_COLUMNS if c not in df.columns]
    return faltantes


# ---------------------------------------------------------------------------
# Procesamiento semanal (diff contra la tabla maestra)
# ---------------------------------------------------------------------------

def procesar_semana(sh, df, semana):
    master_df = load_master(sh)
    master_by_key = (
        master_df.set_index("cliente_key").to_dict("index") if not master_df.empty else {}
    )
    filas_finales = dict(master_by_key)  # copia mutable: clave -> dict de columnas

    enviados = 0
    sin_whatsapp = 0
    derivados = 0
    ventas_total = 0.0
    fecha_hoy = datetime.now().strftime("%Y-%m-%d %H:%M")

    def esta_vacio(x):
        return pd.isna(x) or str(x).strip() == ""

    for _, row in df.iterrows():
        nombre = row.get("Nombre")
        contacto = row.get("Contacto")
        if esta_vacio(nombre) and esta_vacio(contacto):
            continue  # fila vacía, se ignora

        key = build_key(nombre, contacto)
        estado_nuevo = row.get("Estado")
        estado_nuevo = None if pd.isna(estado_nuevo) else str(estado_nuevo).strip()

        venta_nueva = row.get("Venta") if "Venta" in df.columns else None
        venta_nueva = None if pd.isna(venta_nueva) else float(venta_nueva)

        email = row.get("Email") if "Email" in df.columns else None
        localidad = row.get("Localidad") if "Localidad" in df.columns else None
        region = row.get("Región") if "Región" in df.columns else None
        ejecutivo = row.get("Ejecutivo")

        existente = filas_finales.get(key)

        if existente is None:
            if estado_nuevo == ESTADO_ENVIADO:
                enviados += 1
            elif estado_nuevo == ESTADO_SIN_WHATSAPP:
                sin_whatsapp += 1
            elif estado_nuevo == ESTADO_DERIVADO:
                derivados += 1
            if venta_nueva:
                ventas_total += venta_nueva

            filas_finales[key] = {
                "nombre": nombre, "contacto": contacto, "email": email,
                "localidad": localidad, "region": region, "estado": estado_nuevo,
                "ejecutivo": ejecutivo, "venta": venta_nueva,
                "semana_ultimo_cambio": semana, "fecha_actualizacion": fecha_hoy,
            }
        else:
            estado_viejo = existente.get("estado")
            venta_vieja = existente.get("venta")
            venta_vieja = None if pd.isna(venta_vieja) else venta_vieja

            if estado_nuevo != estado_viejo and estado_nuevo is not None:
                if estado_nuevo == ESTADO_ENVIADO:
                    enviados += 1
                elif estado_nuevo == ESTADO_SIN_WHATSAPP:
                    sin_whatsapp += 1
                elif estado_nuevo == ESTADO_DERIVADO:
                    derivados += 1

            if (not venta_vieja) and venta_nueva:
                ventas_total += venta_nueva

            filas_finales[key] = {
                "nombre": nombre, "contacto": contacto, "email": email,
                "localidad": localidad, "region": region,
                "estado": estado_nuevo if estado_nuevo is not None else estado_viejo,
                "ejecutivo": ejecutivo,
                "venta": venta_nueva if venta_nueva else venta_vieja,
                "semana_ultimo_cambio": semana, "fecha_actualizacion": fecha_hoy,
            }

    # Guarda la tabla maestra completa de una sola vez
    if filas_finales:
        nuevo_master = (
            pd.DataFrame.from_dict(filas_finales, orient="index")
            .reset_index()
            .rename(columns={"index": "cliente_key"})
        )
    else:
        nuevo_master = pd.DataFrame(columns=MASTER_COLUMNS)
    guardar_master(sh, nuevo_master)

    # Actualiza (o crea) la fila del historial de esta semana
    hist_df = load_historial(sh)
    if not hist_df.empty and semana in hist_df["semana"].astype(int).values:
        idx = hist_df.index[hist_df["semana"].astype(int) == semana][0]
        hist_df.loc[idx, "whatsapp_enviados"] += enviados
        hist_df.loc[idx, "sin_whatsapp"] += sin_whatsapp
        hist_df.loc[idx, "derivados"] += derivados
        hist_df.loc[idx, "ventas_total"] += ventas_total
        hist_df.loc[idx, "fecha_carga"] = fecha_hoy
    else:
        nueva_fila = pd.DataFrame([{
            "semana": semana, "whatsapp_enviados": enviados,
            "sin_whatsapp": sin_whatsapp, "derivados": derivados,
            "ventas_total": ventas_total, "fecha_carga": fecha_hoy,
        }])
        hist_df = pd.concat([hist_df, nueva_fila], ignore_index=True)

    hist_df = hist_df.sort_values("semana").reset_index(drop=True)
    guardar_historial(sh, hist_df)

    return {
        "enviados": enviados,
        "sin_whatsapp": sin_whatsapp,
        "derivados": derivados,
        "ventas_total": ventas_total,
    }


# ---------------------------------------------------------------------------
# Indicadores adicionales: cierre, distribución geográfica, variación semanal
# ---------------------------------------------------------------------------

def nombre_region(codigo):
    if pd.isna(codigo) or str(codigo).strip() == "":
        return "Sin región"
    codigo_str = str(codigo).strip()
    if codigo_str.endswith(".0"):
        codigo_str = codigo_str[:-2]
    return REGIONES_CHILE.get(codigo_str, f"Región {codigo_str}")


def calcular_metricas_cierre(master_df):
    """Tasa de cierre (derivados que terminan comprando) y ticket promedio,
    calculado sobre el total histórico de clientes (no por semana)."""
    if master_df.empty:
        return {"total_derivados": 0, "total_con_venta": 0, "suma_ventas": 0.0,
                "tasa_cierre": 0.0, "ticket_promedio": 0.0}

    es_derivado = master_df["estado"] == ESTADO_DERIVADO
    tiene_venta = master_df["venta"].notna() & (master_df["venta"] > 0)

    total_derivados = int(es_derivado.sum())
    total_con_venta = int(tiene_venta.sum())
    suma_ventas = float(master_df.loc[tiene_venta, "venta"].sum()) if total_con_venta else 0.0
    tasa_cierre = (total_con_venta / total_derivados * 100) if total_derivados else 0.0
    ticket_promedio = (suma_ventas / total_con_venta) if total_con_venta else 0.0

    return {
        "total_derivados": total_derivados,
        "total_con_venta": total_con_venta,
        "suma_ventas": suma_ventas,
        "tasa_cierre": tasa_cierre,
        "ticket_promedio": ticket_promedio,
    }


def calcular_distribucion_regional(master_df):
    if master_df.empty:
        return pd.DataFrame(columns=["Región", "Derivados", "Clientes con venta", "Ventas ($)"])
    df = master_df.copy()
    df["region_nombre"] = df["region"].apply(nombre_region)
    df["es_derivado"] = df["estado"] == ESTADO_DERIVADO
    df["tiene_venta"] = df["venta"].notna() & (df["venta"] > 0)
    resumen = df.groupby("region_nombre").agg(
        Derivados=("es_derivado", "sum"),
        **{"Clientes con venta": ("tiene_venta", "sum")},
        Ventas=("venta", "sum"),
    ).reset_index().rename(columns={"region_nombre": "Región", "Ventas": "Ventas ($)"})
    return resumen.sort_values("Ventas ($)", ascending=False).reset_index(drop=True)


def calcular_top_localidades(master_df, top_n=10):
    if master_df.empty or "localidad" not in master_df.columns:
        return pd.DataFrame(columns=["Localidad", "Derivados", "Ventas ($)"])
    df = master_df.copy()
    df["es_derivado"] = df["estado"] == ESTADO_DERIVADO
    df["localidad"] = df["localidad"].fillna("Sin localidad")
    resumen = df.groupby("localidad").agg(
        Derivados=("es_derivado", "sum"),
        Ventas=("venta", "sum"),
    ).reset_index().rename(columns={"localidad": "Localidad", "Ventas": "Ventas ($)"})
    return resumen.sort_values(["Ventas ($)", "Derivados"], ascending=False).head(top_n).reset_index(drop=True)


def agregar_variaciones(historial_df):
    """Agrega columnas de variación % semana contra semana anterior."""
    df = historial_df.sort_values("semana").reset_index(drop=True).copy()
    for col_origen, col_nueva in [
        ("whatsapp_enviados", "var_enviados"),
        ("derivados", "var_derivados"),
        ("ventas_total", "var_ventas"),
    ]:
        anterior = df[col_origen].shift(1)
        variacion = ((df[col_origen] - anterior) / anterior.replace(0, pd.NA)) * 100
        df[col_nueva] = variacion.replace([float("inf"), float("-inf")], pd.NA)
    return df


def formatear_variacion(valor):
    if pd.isna(valor):
        return "-"
    flecha = "▲" if valor > 0 else ("▼" if valor < 0 else "▬")
    return f"{flecha} {valor:+.1f}%"


# ---------------------------------------------------------------------------
# Generación de informe Word
# ---------------------------------------------------------------------------

def calcular_comentario_automatico(historial):
    """Calcula un pequeño set de datos/insights a partir del historial.
    Punto de partida para ir agregando más análisis a futuro."""
    total_enviados = int(historial["whatsapp_enviados"].sum())
    total_derivados = int(historial["derivados"].sum())
    total_ventas = float(historial["ventas_total"].sum())
    tasa_conversion = (total_derivados / total_enviados * 100) if total_enviados else 0.0

    semana_mas_derivados = None
    if historial["derivados"].sum() > 0:
        semana_mas_derivados = int(historial.loc[historial["derivados"].idxmax(), "semana"])

    semana_mas_ventas = None
    if historial["ventas_total"].sum() > 0:
        semana_mas_ventas = int(historial.loc[historial["ventas_total"].idxmax(), "semana"])

    return {
        "total_enviados": total_enviados,
        "total_derivados": total_derivados,
        "total_ventas": total_ventas,
        "tasa_conversion": tasa_conversion,
        "semana_mas_derivados": semana_mas_derivados,
        "semana_mas_ventas": semana_mas_ventas,
    }


def generar_grafico_estatico(historial):
    """Genera el gráfico de barras (enviados vs ventas) como imagen PNG
    en memoria, para insertarlo en el Word."""
    fig, ax1 = plt.subplots(figsize=(6.5, 3.3))
    x = historial["semana"].astype(float)

    ax1.bar(x - 0.2, historial["whatsapp_enviados"], width=0.4,
            label="WhatsApp enviados", color="#25D366")
    ax1.set_xlabel("Semana")
    ax1.set_ylabel("N° WhatsApp enviados")
    ax1.set_xticks(list(x))
    ax1.set_xticklabels([str(int(s)) for s in x])

    ax2 = ax1.twinx()
    ax2.bar(x + 0.2, historial["ventas_total"], width=0.4,
            label="Ventas ($)", color="#1F3B4D")
    ax2.set_ylabel("Ventas ($)")

    lineas1, etiquetas1 = ax1.get_legend_handles_labels()
    lineas2, etiquetas2 = ax2.get_legend_handles_labels()
    ax1.legend(lineas1 + lineas2, etiquetas1 + etiquetas2, loc="upper left", fontsize=8)

    fig.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def generar_informe_docx(historial, master_df):
    """Arma el informe Word completo: resumen, tabla, distribución geográfica y gráfico."""
    datos = calcular_comentario_automatico(historial)
    cierre = calcular_metricas_cierre(master_df)
    historial_var = agregar_variaciones(historial)
    ultima = historial_var.iloc[-1]

    doc = Document()
    doc.add_heading("Informe de Avance — Campaña WhatsApp Clientes Potenciales", level=1)
    p = doc.add_paragraph(f"Generado el {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    p.runs[0].italic = True

    doc.add_heading("Resumen general", level=2)
    resumen = (
        f"Durante las {len(historial)} semana(s) registradas se han enviado "
        f"{datos['total_enviados']} mensajes de WhatsApp, de los cuales "
        f"{datos['total_derivados']} clientes respondieron y fueron derivados a un "
        f"vendedor (tasa de conversión de {datos['tasa_conversion']:.1f}%). "
        f"Las ventas acumuladas alcanzan ${datos['total_ventas']:,.0f}."
    )
    doc.add_paragraph(resumen)
    doc.add_paragraph(
        f"Del total histórico de {cierre['total_derivados']} clientes derivados a un vendedor, "
        f"{cierre['total_con_venta']} compraron (tasa de cierre de {cierre['tasa_cierre']:.1f}%), "
        f"con un ticket promedio de ${cierre['ticket_promedio']:,.0f} por venta."
    )
    if not pd.isna(ultima["var_enviados"]) or not pd.isna(ultima["var_derivados"]):
        partes = []
        if not pd.isna(ultima["var_enviados"]):
            partes.append(f"los envíos de WhatsApp variaron {ultima['var_enviados']:+.1f}%")
        if not pd.isna(ultima["var_derivados"]):
            partes.append(f"los derivados variaron {ultima['var_derivados']:+.1f}%")
        doc.add_paragraph(f"Respecto a la semana anterior, {' y '.join(partes)}.")
    if datos["semana_mas_derivados"] is not None:
        doc.add_paragraph(
            f"La semana {datos['semana_mas_derivados']} tuvo la mayor cantidad de clientes derivados."
        )
    if datos["semana_mas_ventas"] is not None:
        doc.add_paragraph(
            f"La semana {datos['semana_mas_ventas']} registró el mayor monto de ventas."
        )

    doc.add_heading("Avance por semana", level=2)
    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = "Light Grid Accent 1"
    encabezados = ["Semana", "WhatsApp enviados", "Derivados", "Ventas ($)"]
    for i, texto in enumerate(encabezados):
        tabla.rows[0].cells[i].text = texto
    for _, fila in historial.iterrows():
        celdas = tabla.add_row().cells
        celdas[0].text = str(int(fila["semana"]))
        celdas[1].text = str(int(fila["whatsapp_enviados"]))
        celdas[2].text = str(int(fila["derivados"]))
        celdas[3].text = f"${fila['ventas_total']:,.0f}"

    doc.add_heading("Distribución geográfica por región", level=2)
    dist_regional = calcular_distribucion_regional(master_df)
    if dist_regional.empty:
        doc.add_paragraph("Sin datos de región disponibles todavía.")
    else:
        tabla_reg = doc.add_table(rows=1, cols=4)
        tabla_reg.style = "Light Grid Accent 1"
        for i, texto in enumerate(["Región", "Derivados", "Clientes con venta", "Ventas ($)"]):
            tabla_reg.rows[0].cells[i].text = texto
        for _, fila in dist_regional.iterrows():
            celdas = tabla_reg.add_row().cells
            celdas[0].text = str(fila["Región"])
            celdas[1].text = str(int(fila["Derivados"]))
            celdas[2].text = str(int(fila["Clientes con venta"]))
            celdas[3].text = f"${fila['Ventas ($)']:,.0f}"

    doc.add_heading("WhatsApp enviados vs Ventas", level=2)
    imagen = generar_grafico_estatico(historial)
    doc.add_picture(imagen, width=Inches(6))

    salida = BytesIO()
    doc.save(salida)
    salida.seek(0)
    return salida


# ---------------------------------------------------------------------------
# UI
# ---------------------------------------------------------------------------

st.set_page_config(page_title="Reporte avance WhatsApp", layout="wide")
st.title("📊 Reporte de avance — Campaña WhatsApp clientes potenciales")

try:
    sh = get_spreadsheet()
except Exception as e:
    st.error(
        "No se pudo conectar con Google Sheets. Revisa que los 'Secrets' "
        "(gcp_service_account y sheet_id) estén bien configurados, y que "
        "la hoja esté compartida con el email de la cuenta de servicio.\n\n"
        f"Detalle técnico: {e}"
    )
    st.stop()

with st.sidebar:
    st.header("Cargar actualización semanal")

    historial_actual = load_historial(sh)
    siguiente_semana = int(historial_actual["semana"].max() + 1) if len(historial_actual) else 1

    semana_num = st.number_input(
        "N° de semana de esta actualización",
        min_value=1, step=1, value=siguiente_semana
    )

    fila_encabezado = st.number_input(
        "Fila donde están los títulos de columna (1 = primera fila)",
        min_value=1, step=1, value=1,
        help="Cámbialo solo si tu Excel tiene una fila de título o logo arriba de Nombre, Contacto, etc."
    )

    archivo = st.file_uploader("Excel de esta semana (.xlsx)", type=["xlsx"])

    if archivo is not None:
        try:
            hojas = pd.ExcelFile(archivo).sheet_names
        except Exception as e:
            st.error(f"No se pudo abrir el archivo: {e}")
            hojas = []

        hoja_seleccionada = None
        if len(hojas) > 1:
            hoja_seleccionada = st.selectbox("Hoja del Excel a usar", hojas)
        elif len(hojas) == 1:
            hoja_seleccionada = hojas[0]

        df_nuevo = None
        columnas_originales = []
        if hoja_seleccionada is not None:
            try:
                df_nuevo = pd.read_excel(
                    archivo, sheet_name=hoja_seleccionada, header=int(fila_encabezado) - 1
                )
                columnas_originales = list(df_nuevo.columns)
                df_nuevo = normalizar_columnas(df_nuevo)
            except Exception as e:
                st.error(f"No se pudo leer el archivo: {e}")
                df_nuevo = None

        if df_nuevo is not None:
            faltantes = validar_columnas(df_nuevo)
            if faltantes:
                st.error(f"Faltan columnas requeridas en el Excel: {', '.join(faltantes)}")
                st.caption("Columnas que se detectaron tal cual en tu archivo:")
                st.code(str(columnas_originales))
                st.caption(
                    "Vista previa cruda de las primeras filas (sin asumir dónde "
                    "están los títulos) — sirve para ubicar en qué fila y hoja "
                    "están realmente Nombre, Contacto, etc.:"
                )
                try:
                    vista_cruda = pd.read_excel(
                        archivo, sheet_name=hoja_seleccionada, header=None, nrows=10
                    )
                    st.dataframe(vista_cruda, use_container_width=True)
                except Exception as e:
                    st.caption(f"No se pudo generar la vista previa: {e}")
            else:
                st.success(f"Archivo leído: {len(df_nuevo)} filas")
                if st.button("Procesar y agregar al historial", type="primary"):
                    resultado = procesar_semana(sh, df_nuevo, int(semana_num))
                    registrar_carga(sh, int(semana_num), archivo.name, len(df_nuevo), resultado)
                    st.success(
                        f"Semana {int(semana_num)} procesada:\n\n"
                        f"- WhatsApp enviados (nuevos): {resultado['enviados']}\n"
                        f"- Sin WhatsApp (nuevos): {resultado['sin_whatsapp']}\n"
                        f"- Derivados (nuevos): {resultado['derivados']}\n"
                        f"- Ventas detectadas: ${resultado['ventas_total']:,.0f}"
                    )
                    st.rerun()

    st.divider()
    with st.expander("⚠️ Zona de pruebas"):
        st.caption("Solo usar mientras pruebas la app — borra todo el historial y la base de clientes.")
        if st.checkbox("Confirmo que quiero borrar todos los datos"):
            if st.button("Reiniciar base de datos"):
                guardar_master(sh, pd.DataFrame(columns=MASTER_COLUMNS))
                guardar_historial(sh, pd.DataFrame(columns=HISTORIAL_COLUMNS))
                st.rerun()

# ---------------------------------------------------------------------------
# Reporte principal
# ---------------------------------------------------------------------------

historial = load_historial(sh)
master_df = load_master(sh)

st.subheader("Avance por semana")

if historial.empty:
    st.info("Aún no hay datos cargados. Sube el primer Excel desde el panel lateral.")
else:
    historial_var = agregar_variaciones(historial)
    ultima = historial_var.iloc[-1]

    col1, col2, col3 = st.columns(3)
    col1.metric(
        "WhatsApp enviados (última semana)", int(ultima["whatsapp_enviados"]),
        delta=None if pd.isna(ultima["var_enviados"]) else f"{ultima['var_enviados']:+.1f}% vs semana anterior",
    )
    col2.metric(
        "Derivados (última semana)", int(ultima["derivados"]),
        delta=None if pd.isna(ultima["var_derivados"]) else f"{ultima['var_derivados']:+.1f}% vs semana anterior",
    )
    col3.metric(
        "Ventas (última semana)", f"${ultima['ventas_total']:,.0f}",
        delta=None if pd.isna(ultima["var_ventas"]) else f"{ultima['var_ventas']:+.1f}% vs semana anterior",
    )

    tabla_mostrar = historial_var.rename(columns={
        "semana": "Semana",
        "whatsapp_enviados": "WhatsApp enviados",
        "sin_whatsapp": "Sin WhatsApp",
        "derivados": "Derivados",
        "ventas_total": "Ventas ($)",
        "fecha_carga": "Última carga",
    })
    tabla_mostrar["Var. enviados"] = historial_var["var_enviados"].apply(formatear_variacion)
    tabla_mostrar["Var. derivados"] = historial_var["var_derivados"].apply(formatear_variacion)
    tabla_mostrar["Var. ventas"] = historial_var["var_ventas"].apply(formatear_variacion)

    st.dataframe(
        tabla_mostrar[["Semana", "WhatsApp enviados", "Var. enviados",
                        "Derivados", "Var. derivados", "Ventas ($)", "Var. ventas",
                        "Sin WhatsApp", "Última carga"]],
        use_container_width=True, hide_index=True
    )

    st.subheader("Indicadores de cierre")
    metricas_cierre = calcular_metricas_cierre(master_df)
    col1, col2, col3 = st.columns(3)
    col1.metric("Total derivados (histórico)", metricas_cierre["total_derivados"])
    col2.metric("Tasa de cierre (derivados → venta)", f"{metricas_cierre['tasa_cierre']:.1f}%")
    col3.metric("Ticket promedio", f"${metricas_cierre['ticket_promedio']:,.0f}")

    st.subheader("Distribución geográfica")
    col_izq, col_der = st.columns(2)
    with col_izq:
        st.caption("Por región")
        st.dataframe(calcular_distribucion_regional(master_df), use_container_width=True, hide_index=True)
    with col_der:
        st.caption("Top 10 localidades")
        st.dataframe(calcular_top_localidades(master_df), use_container_width=True, hide_index=True)

    st.subheader("WhatsApp enviados vs Ventas por semana")

    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=historial["semana"], y=historial["whatsapp_enviados"],
        name="WhatsApp enviados", yaxis="y1", marker_color="#25D366"
    ))
    fig.add_trace(go.Bar(
        x=historial["semana"], y=historial["ventas_total"],
        name="Ventas ($)", yaxis="y2", marker_color="#1F3B4D"
    ))
    fig.update_layout(
        barmode="group",
        xaxis=dict(title="Semana", dtick=1),
        yaxis=dict(title="N° WhatsApp enviados"),
        yaxis2=dict(title="Ventas ($)", overlaying="y", side="right"),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        margin=dict(t=40),
    )
    st.plotly_chart(fig, use_container_width=True)

    st.subheader("Informe descargable")
    informe_docx = generar_informe_docx(historial, master_df)
    st.download_button(
        "📄 Descargar informe en Word",
        data=informe_docx,
        file_name=f"informe_avance_whatsapp_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    with st.expander("📋 Historial de archivos subidos"):
        log_cargas = load_log_cargas(sh)
        if log_cargas.empty:
            st.caption("Todavía no hay cargas registradas.")
        else:
            st.dataframe(
                log_cargas.rename(columns={
                    "fecha_hora": "Fecha y hora", "semana": "Semana",
                    "nombre_archivo": "Archivo", "filas_en_archivo": "Filas en el archivo",
                    "enviados_nuevos": "Enviados (nuevos)", "sin_whatsapp_nuevos": "Sin WhatsApp (nuevos)",
                    "derivados_nuevos": "Derivados (nuevos)", "ventas_detectadas": "Ventas detectadas ($)",
                }),
                use_container_width=True, hide_index=True,
            )

    with st.expander("Ver base de clientes completa (tabla maestra)"):
        st.dataframe(master_df, use_container_width=True, hide_index=True)
